from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Lab-Automation-v2-Operator-Guide.docx"


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 8, (46, 116, 181)),
        ("Heading 2", 13, 12, 6, (46, 116, 181)),
        ("Heading 3", 12, 8, 4, (31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        fmt = style.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = 1.1


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, size=22, bold=True, color=(11, 37, 69))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(subtitle)
    set_font(run2, size=10, color=(80, 80, 80))


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_font(run, name="Consolas", size=9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_two_col_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Topic", bold=True)
    set_cell_text(hdr[1], "Value", bold=True)
    shade_cell(hdr[0], "E8EEF5")
    shade_cell(hdr[1], "E8EEF5")
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True)
        set_cell_text(cells[1], right)


def build() -> Path:
    doc = Document()
    set_page(doc)
    style_doc(doc)
    add_title(
        doc,
        "Lab Automation v2.0 Operator Guide",
        "Camera bridge, live stream checks, AI PC viewing, and live model validation. Built for the current project state on June 16, 2026.",
    )

    doc.add_paragraph(
        "This guide is meant for hands-on testing on the camera Raspberry Pi `hari` and the Windows AI PC. "
        "It stays inside the camera and AI-vision scope: no Auto mode, no relay `/set` commands, and no Node-RED or ESP32 control steps."
    )

    doc.add_heading("1. Current Ground Truth", level=1)
    add_two_col_table(
        doc,
        [
            ("Camera Pi hostname", "hari"),
            ("Expected local bridge URL on hari", "rtsp://127.0.0.1:8554/labcam"),
            ("Expected AI PC URL", "rtsp://hari:8554/labcam"),
            ("Camera IP", "192.168.5.110"),
            ("Working upstream camera path", "rtsps://<user>:<pass>@192.168.5.110:8554/video/live?channel=1&subtype=1&unicast=true&proto=Onvif"),
            ("Known wrong path that returns 404", "rtsps://<user>:<pass>@192.168.5.110:8554/cam/realmonitor?channel=1&subtype=1"),
            ("Production people model", "models/backcam_yolov8s_improved_v3_hardfp.pt"),
            ("Live validator script", "tools/live_model_validation.py"),
        ],
    )

    doc.add_paragraph(
        "Important: a `404 Not Found` from `rtsp://127.0.0.1:8554/labcam` or `rtsp://hari:8554/labcam` usually means MediaMTX is up, but there is no active publisher feeding the `labcam` path. "
        "A `404` from the upstream camera RTSPS URL usually means the camera path itself is wrong."
    )

    doc.add_heading("2. Quick Test Flow", level=1)
    add_numbered(
        doc,
        [
            "On `hari`, test whether the camera IP and ports are reachable.",
            "On `hari`, confirm MediaMTX and the ffmpeg bridge are running.",
            "On `hari`, probe `rtsp://127.0.0.1:8554/labcam`.",
            "On the AI PC, probe `rtsp://hari:8554/labcam`.",
            "If the stream opens, show it in a live window on the AI PC.",
            "If the stream is stable, run `tools/live_model_validation.py` in display mode.",
        ],
    )

    doc.add_heading("3. Commands to Run on hari", level=1)
    doc.add_heading("3.1 Network Reachability", level=2)
    add_code_block(
        doc,
        """
ping -c 4 -W 2 192.168.5.110
nc -vz -w 5 192.168.5.110 80
nc -vz -w 5 192.168.5.110 554
nc -vz -w 5 192.168.5.110 8554
ip addr
ip route
""",
    )

    doc.add_heading("3.2 MediaMTX and Bridge Status", level=2)
    add_code_block(
        doc,
        """
systemctl --user status mediamtx.service
systemctl --user status labos-camera-bridge.service
systemctl --user status labcam-healthcheck.timer
ps -ef | grep -Ei 'mediamtx|ffmpeg|labcam' | grep -v grep
ss -ltnp 2>/dev/null | grep -E ':8554|:8888|:1935|:8889'
grep -n -A8 -B4 labcam /home/hari/mediamtx/mediamtx.yml
""",
    )

    doc.add_heading("3.3 Probe the Upstream Camera Correctly", level=2)
    doc.add_paragraph("Use the working RTSPS path, not the old `/cam/realmonitor` path.")
    add_code_block(
        doc,
        """
ffprobe -hide_banner -v error -rtsp_transport tcp -show_entries stream=codec_name,codec_type,width,height,avg_frame_rate -of json 'rtsps://<user>:<pass>@192.168.5.110:8554/video/live?channel=1&subtype=1&unicast=true&proto=Onvif'
""",
    )

    doc.add_heading("3.4 Probe the Local Published Stream", level=2)
    add_code_block(
        doc,
        """
ffprobe -hide_banner -v error -rtsp_transport tcp -show_streams -show_format -of json rtsp://127.0.0.1:8554/labcam
ffmpeg -hide_banner -loglevel warning -rtsp_transport tcp -i rtsp://127.0.0.1:8554/labcam -t 30 -an -f null -
""",
    )

    doc.add_heading("3.5 If labcam Returns 404", level=2)
    add_bullets(
        doc,
        [
            "If the upstream RTSPS probe fails with 404, the camera path is wrong.",
            "If the upstream RTSPS probe works but `/labcam` is 404, the bridge is stale or absent.",
            "Restart only the bridge first; do not reboot the Pi.",
        ],
    )
    add_code_block(
        doc,
        """
systemctl --user restart labos-camera-bridge.service
sleep 3
ffprobe -hide_banner -v error -rtsp_transport tcp -show_streams -show_format -of json rtsp://127.0.0.1:8554/labcam
""",
    )

    doc.add_heading("4. Commands to Run on the AI PC", level=1)
    doc.add_heading("4.1 Enter the Project", level=2)
    add_code_block(
        doc,
        r"""
cd 'C:\Users\prith\Downloads\Lab Automation v2.0'
""",
    )

    doc.add_heading("4.2 Confirm the Stream Opens from Windows", level=2)
    add_code_block(
        doc,
        r"""
ffprobe -hide_banner -v error -rtsp_transport tcp -show_streams -show_format -of json rtsp://hari:8554/labcam
ffmpeg -hide_banner -loglevel warning -rtsp_transport tcp -i rtsp://hari:8554/labcam -t 30 -an -f null -
""",
    )

    doc.add_heading("4.3 See the Live Footage in a Separate Window", level=2)
    add_code_block(
        doc,
        r"""
ffplay -rtsp_transport tcp -fflags nobuffer -flags low_delay rtsp://hari:8554/labcam
""",
    )
    doc.add_paragraph(
        "That is the quickest plain video-view command. If `ffplay` is unavailable on Windows, the project’s live validator display mode is the fallback viewer."
    )

    doc.add_heading("4.4 Run the Project Live Viewer With Overlay", level=2)
    add_code_block(
        doc,
        r"""
.\.venv\Scripts\python.exe tools\live_model_validation.py --source rtsp://hari:8554/labcam --model models/backcam_yolov8s_improved_v3_hardfp.pt --conf 0.35 --imgsz 1280 --device 0 --tracker bytetrack.yaml --zones config/zones.json --duration 600 --display --output-dir monitor-results\live-validation
""",
    )

    doc.add_paragraph(
        "This opens a live OpenCV window with boxes, per-person confidence, zone labels, total count, per-zone counts, FPS, and inference latency. "
        "It also writes annotated video, CSV, JSONL, and a summary JSON to `monitor-results\\live-validation`."
    )

    doc.add_heading("4.5 Monitor-Safe Vision Publisher Run", level=2)
    add_code_block(
        doc,
        r"""
.\.venv\Scripts\python.exe set_manual_mode.py
.\.venv\Scripts\python.exe -u production_inference_mqtt.py
""",
    )
    add_bullets(
        doc,
        [
            "Use this only when you intentionally want MQTT vision output.",
            "During safe testing, published topics must stay under `labos/v2/vision/#` only.",
            "Monitor mode must send zero relay `/set` commands.",
        ],
    )

    doc.add_heading("5. How to Read the Common Failure Cases", level=1)
    add_two_col_table(
        doc,
        [
            ("`404` on upstream RTSPS URL", "Wrong camera path. Use `/video/live?...&unicast=true&proto=Onvif`."),
            ("`404` on `/labcam` only", "MediaMTX is running but nothing is publishing into `labcam`."),
            ("Port 8554 open but ffprobe fails", "The RTSP server is reachable, but the stream path or publisher state is wrong."),
            ("Port 554 refused but 8554 works", "Expected for the current camera setup; use secure RTSPS on 8554."),
            ("AI PC fails but hari local works", "Name resolution, Windows firewall, or network path issue between AI PC and hari."),
            ("Repeated HEVC warnings during decode", "Usually stream quality or reference-frame issues, not a bridge-path problem by themselves."),
        ],
    )

    doc.add_heading("6. Commands for Zone and Vision Work", level=1)
    add_code_block(
        doc,
        r"""
.\.venv\Scripts\python.exe tools\render_zone_overlay.py --reference monitor-results\zone-calibration\reference.png --zones config\zones.json --output monitor-results\zone-calibration\zone-overlay.png
.\.venv\Scripts\python.exe ai-pc\tools\zone_editor.py --image monitor-results\zone-calibration\reference.png --load config\zones.json --output config\zones.json
.\.venv\Scripts\python.exe tools\validate_zone_geometry.py --zones config\zones.json
.\.venv\Scripts\python.exe -m pytest tests\test_zone_geometry.py tests\test_live_model_validation.py -q
""",
    )

    doc.add_heading("7. Best Practice Structure Going Forward", level=1)
    add_bullets(
        doc,
        [
            "On `hari`, keep only bridge-related assets: MediaMTX config, bridge service, health-check service, and a redacted operator note.",
            "On the AI PC, keep project runtime commands inside this repo so tests, live validation, and model runs stay reproducible.",
            "Treat `labos` separately as the automation broker/controller side; do not mix camera-bridge recovery steps into relay-control docs.",
            "Keep secrets in environment files or local-only config, never in git-tracked scripts or Markdown.",
            "Use one quick-check section and one deep-diagnosis section in future docs so somebody under pressure can recover the stream fast.",
        ],
    )

    doc.add_heading("8. Recommended Daily Checklist", level=1)
    add_numbered(
        doc,
        [
            "Run local `hari` ffprobe on `rtsp://127.0.0.1:8554/labcam`.",
            "Run AI PC ffprobe on `rtsp://hari:8554/labcam`.",
            "Open a short live viewer window.",
            "Run ten-minute live validator with `--display` when doing model checks.",
            "Review `monitor-results` outputs instead of guessing from memory.",
            "Only after all of that, consider any larger deployment or documentation update.",
        ],
    )

    doc.add_heading("9. Single Most Useful Commands", level=1)
    add_two_col_table(
        doc,
        [
            ("Check local stream on hari", "ffprobe -hide_banner -v error -rtsp_transport tcp -show_streams -show_format -of json rtsp://127.0.0.1:8554/labcam"),
            ("Check AI PC stream", "ffprobe -hide_banner -v error -rtsp_transport tcp -show_streams -show_format -of json rtsp://hari:8554/labcam"),
            ("Open live window on AI PC", "ffplay -rtsp_transport tcp -fflags nobuffer -flags low_delay rtsp://hari:8554/labcam"),
            ("Restart stale bridge on hari", "systemctl --user restart labos-camera-bridge.service"),
            ("Run live AI overlay", r".\.venv\Scripts\python.exe tools\live_model_validation.py --source rtsp://hari:8554/labcam --model models/backcam_yolov8s_improved_v3_hardfp.pt --conf 0.35 --imgsz 1280 --device 0 --tracker bytetrack.yaml --zones config/zones.json --duration 600 --display --output-dir monitor-results\live-validation"),
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
